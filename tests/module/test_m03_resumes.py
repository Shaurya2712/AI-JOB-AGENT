import asyncio
from io import BytesIO
from pathlib import Path

from docx import Document
import httpx
from sqlalchemy import select

from app.config import Settings
from app.main import create_app
from app.models.resumes import Resume
from app.schemas.profiles import CandidateProfileInput
from app.services.profiles import ProfileService
from app.services.resume_files import ResumeFileStorage
from app.services.resumes import ResumeService


def build_test_app(database_path: Path, storage_path: Path, *, max_bytes: int = 5 * 1024 * 1024):
    return create_app(
        Settings(
            environment="test",
            database_url=f"sqlite:///{database_path.as_posix()}",
            log_level="WARNING",
            resume_storage_path=storage_path,
            resume_max_bytes=max_bytes,
        )
    )


def make_docx(text: str) -> bytes:
    document = Document()
    document.add_paragraph(text)
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Dart"
    table.cell(0, 1).text = "Flutter"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def make_pdf(text: str) -> bytes:
    escaped_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    content = f"BT\n/F1 12 Tf\n72 720 Td\n({escaped_text}) Tj\nET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(content)).encode("ascii") + b" >>\nstream\n" + content + b"\nendstream",
    ]

    document = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for object_number, value in enumerate(objects, start=1):
        offsets.append(len(document))
        document.extend(f"{object_number} 0 obj\n".encode("ascii"))
        document.extend(value)
        document.extend(b"\nendobj\n")

    xref_offset = len(document)
    document.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    document.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        document.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    document.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(document)


def test_multiple_formats_extract_text_and_primary_selection(tmp_path: Path) -> None:
    storage_path = tmp_path / "controlled-resumes"
    application = build_test_app(tmp_path / "resumes.db", storage_path)

    async def scenario() -> None:
        async with application.router.lifespan_context(application):
            with application.state.session_factory() as session:
                profile = ProfileService(session).create_profile(
                    CandidateProfileInput(
                        name="Mobile",
                        years_experience=5,
                        target_roles=["Flutter Developer"],
                    )
                )
                profile_id = profile.id

            transport = httpx.ASGITransport(app=application)
            async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
                text_upload = await client.post(
                    f"/profiles/{profile_id}/resumes",
                    data={"resume_name": "Plain text"},
                    files={
                        "resume_file": (
                            "../../unsafe-name.txt",
                            b"Flutter engineer with Firebase experience",
                            "text/plain",
                        )
                    },
                )
                with application.state.session_factory() as session:
                    first_resume = session.scalar(select(Resume))
                    assert first_resume is not None
                    assert first_resume.is_primary is True

                docx_upload = await client.post(
                    f"/profiles/{profile_id}/resumes",
                    data={"resume_name": "Document resume"},
                    files={
                        "resume_file": (
                            "resume.docx",
                            make_docx("React Native engineer"),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    },
                )
                with application.state.session_factory() as session:
                    first_two = list(session.scalars(select(Resume).order_by(Resume.id)))
                    assert [resume.is_primary for resume in first_two] == [True, False]

                pdf_upload = await client.post(
                    f"/profiles/{profile_id}/resumes",
                    data={"resume_name": "PDF resume", "make_primary": "on"},
                    files={"resume_file": ("resume.pdf", make_pdf("Python AI Developer"), "application/pdf")},
                )
                page = await client.get("/profiles")

            assert text_upload.status_code == 303
            assert docx_upload.status_code == 303
            assert pdf_upload.status_code == 303
            assert page.status_code == 200
            assert "Plain text" in page.text
            assert "Document resume" in page.text
            assert "PDF resume" in page.text

            with application.state.session_factory() as session:
                resumes = list(session.scalars(select(Resume).order_by(Resume.id)))
                assert len(resumes) == 3
                assert [resume.is_primary for resume in resumes] == [False, False, True]
                assert "Flutter engineer" in resumes[0].extracted_text
                assert "React Native engineer" in resumes[1].extracted_text
                assert "Dart\tFlutter" in resumes[1].extracted_text
                assert "Python AI Developer" in resumes[2].extracted_text
                assert all(Path(resume.file_path).name == resume.file_path for resume in resumes)
                assert all((storage_path / resume.file_path).is_file() for resume in resumes)

                matching_text = ResumeService(
                    session,
                    ResumeFileStorage(storage_path, 5 * 1024 * 1024),
                ).get_extracted_text(profile_id, resumes[0].id)
                assert matching_text == "Flutter engineer with Firebase experience"
                docx_resume_id = resumes[1].id

            transport = httpx.ASGITransport(app=application)
            async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
                primary_response = await client.post(
                    f"/profiles/{profile_id}/resumes/{docx_resume_id}/primary"
                )
                repeated_primary_response = await client.post(
                    f"/profiles/{profile_id}/resumes/{docx_resume_id}/primary"
                )

            assert primary_response.status_code == 303
            assert repeated_primary_response.status_code == 303
            with application.state.session_factory() as session:
                resumes = list(session.scalars(select(Resume).order_by(Resume.id)))
                assert [resume.is_primary for resume in resumes] == [False, True, False]

    asyncio.run(scenario())


def test_unsupported_and_oversized_uploads_are_not_persisted(tmp_path: Path) -> None:
    storage_path = tmp_path / "controlled-resumes"
    application = build_test_app(tmp_path / "resumes.db", storage_path, max_bytes=64)

    async def scenario() -> None:
        async with application.router.lifespan_context(application):
            with application.state.session_factory() as session:
                profile = ProfileService(session).create_profile(
                    CandidateProfileInput(name="Mobile", target_roles=["Flutter Developer"])
                )
                profile_id = profile.id

            transport = httpx.ASGITransport(app=application)
            async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
                unsupported = await client.post(
                    f"/profiles/{profile_id}/resumes",
                    data={"resume_name": "Executable"},
                    files={"resume_file": ("resume.exe", b"not executable", "application/octet-stream")},
                )
                oversized = await client.post(
                    f"/profiles/{profile_id}/resumes",
                    data={"resume_name": "Too large"},
                    files={"resume_file": ("resume.txt", b"x" * 65, "text/plain")},
                )

            assert unsupported.status_code == 422
            assert "TXT, PDF, or DOCX" in unsupported.text
            assert oversized.status_code == 422
            assert "upload limit" in oversized.text

            with application.state.session_factory() as session:
                assert list(session.scalars(select(Resume))) == []
            assert not storage_path.exists() or not any(storage_path.iterdir())

    asyncio.run(scenario())
