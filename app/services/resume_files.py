from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from uuid import uuid4
from zipfile import BadZipFile, ZipFile

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}
MAX_PDF_PAGES = 100
MAX_DOCX_FILES = 500
MAX_DOCX_UNCOMPRESSED_BYTES = 25 * 1024 * 1024
MAX_EXTRACTED_CHARACTERS = 1_000_000


class ResumeFileError(ValueError):
    pass


@dataclass(frozen=True)
class StoredResume:
    storage_reference: str
    extracted_text: str


class ResumeFileStorage:
    def __init__(self, root: Path, max_bytes: int) -> None:
        self.root = root.expanduser().resolve()
        self.max_bytes = max_bytes

    def store(self, original_filename: str, content: bytes) -> StoredResume:
        extension = Path(original_filename).suffix.casefold()
        if extension not in SUPPORTED_EXTENSIONS:
            raise ResumeFileError("Resume must be a TXT, PDF, or DOCX file")
        if not content:
            raise ResumeFileError("Resume file is empty")
        if len(content) > self.max_bytes:
            raise ResumeFileError(f"Resume exceeds the {self.max_bytes // (1024 * 1024)} MiB upload limit")

        extracted_text = self._extract(extension, content)
        if not extracted_text:
            raise ResumeFileError("Resume contains no readable text")
        if len(extracted_text) > MAX_EXTRACTED_CHARACTERS:
            raise ResumeFileError("Extracted resume text is too large")

        self.root.mkdir(parents=True, exist_ok=True)
        storage_reference = f"{uuid4().hex}{extension}"
        destination = self._resolve(storage_reference)
        destination.write_bytes(content)
        return StoredResume(storage_reference=storage_reference, extracted_text=extracted_text)

    def remove(self, storage_reference: str) -> None:
        destination = self._resolve(storage_reference)
        destination.unlink(missing_ok=True)

    def _resolve(self, storage_reference: str) -> Path:
        if Path(storage_reference).name != storage_reference:
            raise ResumeFileError("Invalid resume storage reference")
        destination = (self.root / storage_reference).resolve()
        if destination.parent != self.root:
            raise ResumeFileError("Resume storage path is outside the configured directory")
        return destination

    def _extract(self, extension: str, content: bytes) -> str:
        if extension == ".txt":
            text = self._extract_txt(content)
        elif extension == ".pdf":
            text = self._extract_pdf(content)
        else:
            text = self._extract_docx(content)
        return "\n".join(line.rstrip() for line in text.splitlines()).strip()

    @staticmethod
    def _extract_txt(content: bytes) -> str:
        try:
            return content.decode("utf-8-sig")
        except UnicodeDecodeError as error:
            raise ResumeFileError("TXT resumes must use UTF-8 encoding") from error

    @staticmethod
    def _extract_pdf(content: bytes) -> str:
        if not content.startswith(b"%PDF-"):
            raise ResumeFileError("The uploaded file is not a valid PDF")
        try:
            reader = PdfReader(BytesIO(content), strict=False)
            if reader.is_encrypted:
                raise ResumeFileError("Encrypted PDF resumes are not supported")
            if len(reader.pages) > MAX_PDF_PAGES:
                raise ResumeFileError(f"PDF resumes may contain at most {MAX_PDF_PAGES} pages")
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except ResumeFileError:
            raise
        except Exception as error:
            raise ResumeFileError("The PDF could not be read") from error

    @staticmethod
    def _extract_docx(content: bytes) -> str:
        try:
            with ZipFile(BytesIO(content)) as archive:
                entries = archive.infolist()
                if len(entries) > MAX_DOCX_FILES:
                    raise ResumeFileError("DOCX archive contains too many files")
                if sum(entry.file_size for entry in entries) > MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise ResumeFileError("DOCX archive expands beyond the safe size limit")
            document = Document(BytesIO(content))
        except ResumeFileError:
            raise
        except (BadZipFile, ValueError, KeyError) as error:
            raise ResumeFileError("The uploaded file is not a valid DOCX") from error
        except Exception as error:
            raise ResumeFileError("The DOCX could not be read") from error

        values = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                values.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(values)
